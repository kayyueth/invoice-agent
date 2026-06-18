"""Fill the Word template and export PDF."""
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table


def _replace_in_paragraph(p, mapping):
    # python-docx splits runs unpredictably; join, replace, write back.
    full = "".join(r.text for r in p.runs)
    if not any(k in full for k in mapping):
        return
    for k, v in mapping.items():
        full = full.replace(k, str(v))
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = full
    else:
        p.add_run(full)


def _all_tables(doc):
    """All tables including nested ones."""
    return [Table(tbl, doc.part) for tbl in doc.element.body.findall('.//' + qn('w:tbl'))]


def _replace_everywhere(doc, mapping):
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)
    for t in _all_tables(doc):
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, mapping)


def _find_detail_table(doc):
    """Detail table = the one whose header row contains 'Date' and 'Category'."""
    for t in _all_tables(doc):
        if not t.rows:
            continue
        header = " ".join(c.text for c in t.rows[0].cells).lower()
        if "date" in header and "category" in header:
            return t
    return None


def _add_detail_rows(table, items):
    """Items: list of dicts with date, unit, category, work_purpose, amount, currency, payment_method, description.
    Template has one sample row after the header; we clone its formatting."""
    if len(table.rows) < 2:
        raise RuntimeError("Detail table needs a sample row to clone formatting from")
    sample = table.rows[1]
    sample_tr = sample._tr
    # Remove sample row; we'll append all real rows with its formatting.
    sample_tr.getparent().remove(sample_tr)

    for it in items:
        new_tr = deepcopy(sample_tr)
        table._tbl.append(new_tr)
        cells = table.rows[-1].cells
        values = [
            it.get("date", ""),
            str(it.get("unit", 1)),
            it.get("category", ""),
            it.get("work_purpose", ""),
            f"{it['amount']:.2f}" if isinstance(it.get("amount"), (int, float)) else str(it.get("amount", "")),
            it.get("payment_method", ""),
            it.get("description", ""),
        ]
        for cell, val in zip(cells, values):
            # Wipe existing paragraph text, keep first paragraph's format.
            if cell.paragraphs:
                p = cell.paragraphs[0]
                for r in p.runs:
                    r.text = ""
                if p.runs:
                    p.runs[0].text = val
                else:
                    p.add_run(val)
                # Drop extra paragraphs
                for extra in cell.paragraphs[1:]:
                    extra._element.getparent().remove(extra._element)


def render(template_path: Path, out_docx: Path, mapping: dict, detail_items: list):
    doc = Document(str(template_path))
    _replace_everywhere(doc, mapping)
    table = _find_detail_table(doc)
    if table is None:
        raise RuntimeError("Could not locate detail table in template")
    _add_detail_rows(table, detail_items)
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))


